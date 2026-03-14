from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import json
import re
from typing import Dict, Any

class Document:
    def __init__(self, path, config_manager=None):
        self.path = path
        self.doc = DocxDocument(path)
        self.config_manager = config_manager
        # 存储论文的各部分内容
        self.title = None
        self.abstract = None
        self.keywords = None
        self.sections = {}     # 格式: {"标题文本": [段落1, 段落2...]}
        self.section_levels = {} # 格式: {"标题文本": 级别数字(1为一级, 2为二级)}
        self.ai_assistant = None

        # 先尝试通过文档样式来解析
        if not self._parse_by_styles():
            # 如果样式解析失败，试试使用传统方法
            if not self._parse_document_traditional():
                # 如果传统方法也失败，且AI功能已启用，则使用AI分析
                if self.config_manager and self.config_manager.is_ai_enabled():
                    from .ai_assistant import DocumentAI
                    self.ai_assistant = DocumentAI(self.config_manager)
                    self._parse_with_ai()

    def _parse_by_styles(self) -> bool:
        """
        通过文档现有样式解析文档结构
        返回是否解析成功
        """
        try:
            current_section = None

            for para in self.doc.paragraphs:
                if not para.text.strip():
                    continue

                style_name = para.style.name.lower() if para.style else ""

                # 通过样式名称识别各部分
                if 'title' in style_name or '标题' in style_name and not any(str(i) in style_name for i in range(1,9)):
                    if not self.title:  # 只取第一个文档标题
                        self.title = para
                elif 'abstract' in style_name or '摘要' in style_name:
                    self.abstract = para
                elif 'keywords' in style_name or '关键词' in style_name:
                    self.keywords = para
                elif 'heading 1' in style_name or '标题 1' in style_name:
                    current_section = para.text.strip()
                    self.sections[current_section] = []
                    self.section_levels[current_section] = 1
                elif 'heading 2' in style_name or '标题 2' in style_name:
                    current_section = para.text.strip()
                    self.sections[current_section] = []
                    self.section_levels[current_section] = 2
                elif current_section:
                    self.sections[current_section].append(para)

            # 如果至少识别出标题和一个章节，则认为解析成功
            return bool(self.title and self.sections)
        except Exception as e:
            print(f"样式解析出错: {str(e)}")
            return False

    def _parse_document_traditional(self) -> bool:
        """
        使用正则和状态机辅助的传统方法解析文档结构
        """
        try:
            current_section = None
            found_structure = False

            # 状态机：记录当前在解析哪一部分 ("title_check", "abstract", "keywords", "body")
            current_state = "title_check"

            paragraphs = self.doc.paragraphs

            for i, para in enumerate(paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                # ------ 边界探测：摘要和关键词 ------
                lower_text = text.lower()
                if lower_text.startswith('abstract') or text.startswith('摘要'):
                    self.abstract = para
                    current_state = "abstract"
                    found_structure = True
                    continue

                if lower_text.startswith('keyword') or text.startswith('关键字') or text.startswith('关键词'):
                    self.keywords = para
                    current_state = "keywords"
                    found_structure = True
                    continue

                # ------ 边界探测：章节标题 ------
                heading_level = self._detect_heading_level(text)
                if heading_level > 0:
                    current_section = text
                    self.sections[current_section] = []
                    self.section_levels[current_section] = heading_level
                    current_state = "body"
                    found_structure = True
                    continue

                # ------ 第一段智能识别（如果既不是摘要也不是章节标题） ------
                if current_state == "title_check" and not self.title:
                    # 如果是很长的一段话（>50字），或者是带句号的句子，极大概率是直接开始写正文/摘要了，跳过标题
                    if len(text) > 50 or text.endswith('。') or text.endswith('.'):
                        current_state = "body" # 放弃抓取文档大标题
                    else:
                        self.title = para
                        found_structure = True
                    continue

                # ------ 内容收集 ------
                if current_state == "abstract" and not self.abstract:
                    pass # 等待
                elif current_state == "keywords" and not self.keywords:
                    pass # 等待
                elif current_section:
                    self.sections[current_section].append(para)

            return found_structure
        except Exception as e:
            print(f"传统解析方法出错: {str(e)}")
            return False

    def _detect_heading_level(self, text: str) -> int:
        """
        通过正则表达式智能判断标题级别
        返回: 1(一级), 2(二级), 0(不是标题)
        """
        text = text.strip()

        # === 一级标题规则 ===

        # 1. 纯数字 + 点 + 文字 (如 "1. 引言" 或 "1 获取数据")，注意这里明确不要匹配第二个点
        if re.match(r'^\d+\s*\.\s*[\u4e00-\u9fa5a-zA-Z]+', text) or re.match(r'^\d+\s+[\u4e00-\u9fa5a-zA-Z]+', text):
            return 1

        # 2. 中文编号 (如 "一、引言", "第一章 概述")
        if re.match(r'^第?[一二三四五六七八九十]+[、\s]\s*[\u4e00-\u9fa5a-zA-Z]+', text):
            return 1

        # 3. 常见大章节独立关键词
        main_keywords = r'^(引言|前言|概论|概述|研究背景|理论基础|研究方法|实验方法|实验|结果分析|实验结果|结果|讨论|结论|总结|参考文献|致谢|附录)$'
        if re.match(main_keywords, text):
            return 1

        # === 二级标题规则 ===

        # 1. 多级数字编号 (如 "1.1 研究背景", "2.1.3 具体实施")
        if re.match(r'^\d+\.\d+(\.\d+)?\s*[\u4e00-\u9fa5a-zA-Z]*', text):
            return 2

        # 2. 带括号的中文编号 (如 "(一) 研究背景", "（二）相关工作")
        if re.match(r'^[(（][一二三四五六七八九十]+[)）]\s*[\u4e00-\u9fa5a-zA-Z]*', text):
            return 2

        return 0

    def _parse_with_ai(self):
        """
        使用AI辅助解析文档结构
        """
        if not self.ai_assistant:
            return False
        
        full_text = "\n".join([para.text for para in self.doc.paragraphs])
        ai_analysis = self.ai_assistant.analyze_document(full_text)
        
        if ai_analysis:
            return self._update_structure_from_ai(ai_analysis)
        return False

    def _update_structure_from_ai(self, ai_analysis: Dict[str, Any]) -> bool:
        """
        根据AI分析结果更新文档结构
        Args:
            ai_analysis: AI分析的结果
        Returns:
            是否成功更新结构
        """
        try:
            # 更新文档各部分
            for para in self.doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 根据AI识别结果匹配段落
                if text == ai_analysis.get('title'):
                    self.title = para
                elif text == ai_analysis.get('abstract'):
                    self.abstract = para
                elif text == ai_analysis.get('keywords'):
                    self.keywords = para
                
                # 处理章节
                for section in ai_analysis.get('sections', []):
                    if text == section['title']:
                        current_section = text
                        self.sections[current_section] = []
                    elif current_section:
                        self.sections[current_section].append(para)
            
            return bool(self.title or self.abstract or self.sections)
        except Exception as e:
            print(f"更新文档结构失败: {str(e)}")
            return False

    def _is_section_heading(self, text: str) -> bool:
        """
        判断是否为章节标题
        """
        # 检查数字编号格式
        if any(text.startswith(f"{i}.") for i in range(1, 10)):
            return True
        
        # 检查中文数字编号格式
        chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
        if any(text.startswith(f"{num}、") for num in chinese_numbers):
            return True
        
        # 检查特定的标题关键词
        heading_keywords = ['引言', '介绍', '研究方法', '实验', '结果', '讨论', '结论', '参考文献']
        return any(keyword in text for keyword in heading_keywords)

    def get_title(self):
        """获取论文标题"""
        return self.title

    def get_abstract(self):
        """获取摘要部分"""
        return self.abstract

    def get_keywords(self):
        """获取关键词部分"""
        return self.keywords

    def get_section(self, section_name):
        """获取指定章节的内容"""
        return self.sections.get(section_name, [])

    def get_all_sections(self):
        """获取所有章节"""
        return self.sections

    def get_references(self):
        """获取参考文献部分"""
        references = []
        for section_name, paragraphs in self.sections.items():
            if '参考文献' in section_name or 'references' in section_name.lower():
                references.extend(paragraphs)
        return references

    def get_tables(self):
        """获取所有表格"""
        return self.doc.tables

    def save(self, output_path: str):
        """
        保存文档
        Args:
            output_path: 输出文件路径
        """
        try:
            self.doc.save(output_path)
            print(f"文档已保存至: {output_path}")
        except Exception as e:
            print(f"保存文档失败: {str(e)}")

    def get_paragraphs(self):
        """获取所有段落"""
        return self.doc.paragraphs

    def get_ai_format_suggestions(self, section_type):
        """获取AI对特定部分的格式建议
        """
        content = None
        if section_type == 'title':
            content = self.title.text if self.title else None
        elif section_type == 'abstract':
            content = self.abstract.text if self.abstract else None
        elif section_type == 'keywords':
            content = self.keywords.text if self.keywords else None
        # 还可以添加其他部分的相关功能
        
        if content:
            return self.ai_assistant.suggest_formatting(section_type, content)
        return None

    def add_section_breaks(self):
        """
        为每个主要章节添加分节符
        """
        try:
            # 获取所有段落
            paragraphs = self.doc.paragraphs
            
            # 遍历段落，为每个一级标题前添加分节符
            for i, para in enumerate(paragraphs):
                if self._is_main_section_heading(para.text):
                    # 在当前段落前添加分节符
                    run = para._p.get_or_add_pPr()
                    sectPr = run.get_or_add_sectPr()
                    # 设置分节类型为下一页
                    sectPr.set('type', 'nextPage')
                    print(f"已在章节 '{para.text}' 前添加分节符")
        
        except Exception as e:
            print(f"添加分节符时出错: {str(e)}")

    def _is_main_section_heading(self, text: str) -> bool:
        """
        判断是否为一级章节标题
        """
        text = text.strip()
        
        # 检查数字编号格式（如 "1. 引言"）
        if any(text.startswith(f"{i}. ") for i in range(1, 10)):
            return True
        
        # 检查中文数字编号格式（如 "一、引言"）
        chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
        if any(text.startswith(f"{num}、") for num in chinese_numbers):
            return True
        
        # 检查特定的一级标题关键词
        main_section_keywords = [
            '引言', '介绍',
            '研究背景', '理论基础',
            '研究方法', '实验方法',
            '结果分析', '实验结果',
            '讨论', '结论',
            '参考文献'
        ]
        
        return any(text.startswith(keyword) for keyword in main_section_keywords)